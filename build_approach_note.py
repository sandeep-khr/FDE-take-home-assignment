from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).parent / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
OUT_FILE = OUT_DIR / "Approach-Note-Flent-FDE.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.3, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_para(doc, text="", *, style=None, before=0, after=5, line=1.1, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def add_lead_para(doc, lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(lead)
    set_run_font(r, size=10.3, color=INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.3, color=INK)
    return p


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_run_font(r, size=10.0, color=INK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.06
    r = p.add_run(text)
    set_run_font(r, size=10.1, color=INK)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(9)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(11.5)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.0)


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FLENT FDE TAKE-HOME  /  MARKET LISTING TRUST LAYER")
    set_run_font(r, size=8.2, color=MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    r = p.add_run("Approach note  |  Problem 1")
    set_run_font(r, size=8.2, color=MUTED)
    p.add_run("\t")
    r = p.add_run("Page ")
    set_run_font(r, size=8.2, color=MUTED)
    add_page_field(p, "PAGE")
    r = p.add_run(" of ")
    set_run_font(r, size=8.2, color=MUTED)
    add_page_field(p, "NUMPAGES")


def add_evidence_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    set_cell_shading(header[0], LIGHT_BLUE)
    set_cell_shading(header[1], LIGHT_BLUE)
    set_cell_text(header[0], "Evidence condition", bold=True, color=DARK_BLUE, size=9.2)
    set_cell_text(header[1], "How I treat it", bold=True, color=DARK_BLUE, size=9.2)
    mark_header_row(table.rows[0])
    rows = [
        ("Missing", "Keep it blank. A missing area weakens a match; it never becomes a made-up number."),
        ("Stale", "More than 30 days dark is excluded. Fifteen to 30 days is kept at half weight."),
        ("Conflicting", "Show the spread inside duplicate clusters and send cross-society near-matches to review."),
        ("Unreliable", "Quarantine or downweight with a written reason. Keep the original row in the audit trail."),
    ]
    for condition, treatment in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], condition, bold=True, color=INK)
        set_cell_text(cells[1], treatment, color=INK)
    set_table_geometry(table, [1800, 7560])
    return table


def add_assumption_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Choice", "Working rule", "Why I chose it")):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True, color=DARK_BLUE, size=9.0)
    mark_header_row(table.rows[0])
    rows = [
        ("Staleness", ">30 days out; 15-30 days at 0.5 weight", "A half-weight middle ground is safer than a sharp cut-off."),
        ("Duplicates", "Dates within 3 days; area within 25 sqft; rent or deposit agrees", "Strict rules protect against merging different homes that merely look alike."),
        ("Aliases", "Two independent cross-posted homes bridge name spellings", "One bridge can itself be a bad listing."),
        ("Confidence", "Effective sample, source count, recency and leave-one-out swing", "A stable number from very thin evidence should still be low confidence."),
    ]
    for choice, rule, why in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], choice, bold=True, size=8.8)
        set_cell_text(cells[1], rule, size=8.8)
        set_cell_text(cells[2], why, size=8.8)
    set_table_geometry(table, [1350, 3650, 4360])
    return table


def add_hyperlink(paragraph, url, text, *, size=9.3, color=BLUE):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    run = paragraph.runs[-1] if paragraph.runs else None
    from docx.text.run import Run
    run = Run(new_run, paragraph)
    set_run_font(run, size=size, color=color, bold=False)
    return run


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.83)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header_footer(section)

    # Page 1: framing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("APPROACH NOTE")
    set_run_font(r, size=9.2, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Problem 1: The market data lies to us")
    set_run_font(r, size=22, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Flent Forward Deployed Engineer take-home | Bangalore rental listings")
    set_run_font(r, size=10.2, color=MUTED)
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Sandeep Kumar  ·  31 August 2026")
    set_run_font(r, size=9.6, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Live demo: ")
    set_run_font(r, size=9.3, color=MUTED)
    add_hyperlink(p, "https://flent-fde-assignment.vercel.app", "flent-fde-assignment.vercel.app")
    r = p.add_run("   ·   Code: ")
    set_run_font(r, size=9.3, color=MUTED)
    add_hyperlink(p, "https://github.com/sandeep-khr/Flent-FDE-take-home-assignment", "github.com/sandeep-khr/Flent-FDE-take-home-assignment")
    r = p.add_run("   ·   Plan board: ")
    set_run_font(r, size=9.3, color=MUTED)
    add_hyperlink(p, "https://github.com/users/sandeep-khr/projects/1/views/2", "github.com/users/sandeep-khr/projects/1/views/2")

    add_heading(doc, "The problem, as I see it", 1)
    add_lead_para(doc, "A bad median looks just as precise as a good one. ",
                  "That is the real failure here. In the case packet, a median across all 86 scraped listings is ₹59,000. It looks clean on a dashboard. But that pull also contains an obvious typo at ₹12,000, an outlier at ₹1,85,000, old listings, broker copies, and the subject home itself posted under the wrong BHK. A number can be neat while the evidence behind it is not.")
    add_para(doc, "The person who feels this is the acquisition owner. They are deciding whether to put a multi-year rent obligation and roughly ₹2.6 lakh of capital behind one home. Supply is also negotiating against the benchmark. If the comparable set is really four homes dressed up as 30, one bad listing can move the answer by thousands. The brief makes the cost concrete: roughly ten extra vacancy days can push breakeven back by a month.")

    add_heading(doc, "The framing I used", 2)
    add_para(doc, "I did not treat this as a data-cleaning exercise. The problem is silent judgment. A system can silently trust a row, or silently delete it, and both choices turn an opinion into arithmetic. My aim is to make those choices visible. Each listing keeps its original values, its grade, and a plain reason for what happened to it. The market estimate then carries the confidence it has actually earned.")
    add_para(doc, "That changes the product from “here is the market rent” to “here is the evidence we have, what it supports, and what is still uncertain.” It is a small distinction in wording, but it changes how safely the team can act.")

    add_heading(doc, "The claim I am willing to make", 2)
    add_para(doc, "For this deal, the layer can recover a defensible asking-rent benchmark from a polluted pull without hiding the messy parts. It should not pretend to predict the rent a tenant will finally sign for. That needs signed-rent history, which this packet does not contain.")

    doc.add_page_break()

    # Page 2: system
    add_heading(doc, "How the trust layer works", 1)
    add_lead_para(doc, "The flow is simple on purpose. ",
                  "Parse the raw file, normalize dates and names, fold obvious broker cross-posts into one physical unit, merge society spellings only when the evidence proves they refer to the same place, grade each row, and calculate a weighted median with a confidence level. Every threshold sits in one config object and is shown in the review surface.")

    add_heading(doc, "What stays deterministic", 2)
    add_para(doc, "Anything that moves the benchmark stays in plain code: parsing, matching rules, weights, median calculation, confidence checks and the final comparison. The same file should give the same result later. That matters when a reviewer asks why a listing was excluded or when a decision needs to be replayed after the outcome is known.")

    add_heading(doc, "Where a person should decide", 2)
    add_para(doc, "I deliberately leave two decisions open for review. First, cross-society near-matches go to a suspect queue rather than being auto-merged. CP-0026 and CP-0053 have similar rents, areas and dates, but their deposits and society names disagree. They may be similar homes, not the same home. Second, a reviewer can exclude or reinstate any row with a required reason. The estimate recomputes and the override sits beside the machine's original reasoning.")

    add_heading(doc, "Where AI helps", 2)
    add_para(doc, "AI is useful upstream: reading messy listing text, proposing normalized fields, and surfacing possible name matches for a person to inspect. I would not put a model inside the final calculation. A fluent explanation that misquotes one rupee is worse than no explanation. The decision math should remain inspectable and deterministic; AI can help prepare evidence, not quietly decide it.")

    add_heading(doc, "How the layer handles imperfect evidence", 2)
    add_evidence_table(doc)

    doc.add_page_break()

    # Page 3: proof and scope
    add_heading(doc, "What the proof shows", 1)
    add_lead_para(doc, "The riskiest assumption was that understandable rules would leave enough good evidence to act on. ",
                  "The end-to-end run supports that claim for this case. The byte-checked CSV goes in unchanged. Twenty-one of the 86 listings contribute to the same-society, 2BHK, semi-furnished comparison set. Their effective sample is 16.0 after downweighting weaker evidence. The weighted-median asking-rent benchmark is ₹59,500, with a ₹58,500 to ₹60,000 bootstrap band and a ₹55,500 to ₹63,000 observed range.")
    add_para(doc, "The result is HIGH confidence under the defined ladder: the evidence comes from four platforms, has a median age of six days, and the leave-one-out swing is ₹0. Removing any one contributor does not change the median. The raw ₹59,000 result was close, but only by accident. In this packet the bad rows happen to offset one another. Cleaning still bought a known sample, a usable range, and evidence that is stable enough to challenge. The same pipeline handles Flent's real volume: a labeled synthetic 27,000-row file runs end to end in under three seconds in the browser, and the page's upload box will take your own pull.")
    add_para(doc, "An early version of my duplicate rule was too loose. It chained 31 different homes into one cluster and cut the comparison set to two rows, so the shipped rule is strict and its near-misses go to a human instead.")

    add_heading(doc, "The trap that mattered most", 2)
    add_para(doc, "CP-0081 is the subject flat cross-posted by a broker with a 3BHK label. It shares the subject's 1,175 sqft area, ₹2.8 lakh deposit and recent posting date. The layer quarantines it as a subject echo. If it were allowed into the comps, the landlord's own ask would help validate itself. It also explains why the 3BHK segment correctly returns N=0 instead of pretending a mislabeled row is evidence.")

    add_heading(doc, "A result can be useful and still be incomplete", 2)
    add_para(doc, "The landlord asks ₹56,000 plus ₹5,000 maintenance. Listing rents do not reliably say whether maintenance is included, so the honest verdict is two readings: the all-in ask is 2.5% above the benchmark; the base rent is 5.9% below it. I would not pick the reading that makes the deal look better. The sign flip is the finding, and it identifies the next piece of data to collect.")

    add_heading(doc, "The deliberate failure case and scope cuts", 2)
    add_para(doc, "When I ask for a furnished benchmark, only four rows survive, with an effective sample of 3.0. The system labels that LOW confidence and gives a collection list rather than a recommendation. For 3BHK it says INSUFFICIENT. This is intentional: refusing a weak answer is part of the product.")
    add_bullet(doc, "I did not build an achieved-rent model. Every listing is an ask; signed rents are needed before it can be calibrated.")
    add_bullet(doc, "I did not impute maintenance, adjust for floor or view, or add synthetic rows. At this sample size, those would create invented precision.")
    add_bullet(doc, "I did not build a locality map. The packet has anonymised societies and no coordinates; the real signal here is behavioural, not geographic.")
    add_bullet(doc, "I kept probabilistic record linkage as the scale path, not the v1 default. Strict matching plus a human queue is easier to audit at 86 rows.")

    doc.add_page_break()

    # Page 4: validation, assumptions, delivery
    add_heading(doc, "What I would validate next", 1)
    add_lead_para(doc, "First: capture whether maintenance is included in each listing rent. ",
                  "It changes this deal's verdict from above market to below market. I would test the field on ten live deals by calling the surviving fresh comps. If at least 80% can state inclusion clearly, we should see the two-reading verdict collapse into one. That is stronger evidence than choosing a default today.")
    add_para(doc, "Second, I would log the benchmark beside eventual signed rent and days-to-fill. A fast delisting may hint that a listing cleared near its ask, but it can also be an expiry or repost, so it remains display-only for now. After enough signed outcomes, the team can measure the gap between asking rents and achieved rents instead of assuming they are the same.")

    add_heading(doc, "Assumptions I made visible", 2)
    add_assumption_table(doc)
    add_para(doc, "These are working assumptions, not Flent policy. They are deliberately round, visible and easy to change. The release should show what changes when a threshold changes, rather than make the constants disappear into the pipeline.")

    add_heading(doc, "How I would take this from proof to use", 2)
    add_numbered(doc, "Start with canonical ingest and the tested rules engine, so every raw field and decision can be replayed.")
    add_numbered(doc, "Add deduplication and alias review with a precision check on sampled clusters before allowing automatic folding at scale.")
    add_numbered(doc, "Ship the benchmark, band and confidence ladder inside the BOSS deal page, with overrides and reasons visible in two clicks.")
    add_numbered(doc, "Run in shadow mode on ten live deals. Watch reviewer override rate and the share of LOW or INSUFFICIENT outputs before widening rollout.")

    add_heading(doc, "How I used AI", 2)
    add_para(doc, "I used Claude to explore the case packet, surface traps, pair-build the tested pipeline, and draft this note from decisions I had already made, which I then rewrote. I checked results through the tests, golden run and independent calculations. I made the choices on framing, thresholds, human review and scope. The repository includes the full AI usage log.", after=0)

    doc.core_properties.title = "Approach Note - Flent FDE Take-home"
    doc.core_properties.subject = "Problem 1: Market Listing Trust Layer"
    doc.core_properties.author = "Sandeep Kumar"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_document()
